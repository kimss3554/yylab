#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os, pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BIRD_PATH  = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아\배이아_새소리.xlsx"
CHART_DIR  = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아\charts"
SRC_DOCX   = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아\배이아_최종탐구보고서_수상권버전.docx"
OUT_DOCX   = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아\배이아_최종탐구보고서_최종제출본.docx"

plt.rcParams['font.family'] = 'Malgun Gothic'
plt.rcParams['axes.unicode_minus'] = False
GREEN_DARK = '#1b4332'; RED = '#c0392b'; GRAY = '#6c757d'; BG = '#f8f9fa'

# ── BirdNET 재현성 분석 ─────────────────────────────────────────────
bird = pd.read_excel(BIRD_PATH, header=1)
cols = list(bird.columns)

# BirdNET: 1차 1순위=col4, 2차 1순위=col7, 3차 1순위=col10
b1 = bird.iloc[:, 4].astype(str).str.strip()
b2 = bird.iloc[:, 7].astype(str).str.strip()
b3 = bird.iloc[:, 10].astype(str).str.strip()

# 지빠귀류 판별
thrush_kw = ['지빠귀', '지빡', 'Thrush', 'thrush', 'Robin', 'robin',
             'Turdus', 'turdus', '울새', '딱새', '개똥지빠귀', '흰배지빠귀']

def is_thrush(s):
    return any(k in str(s) for k in thrush_kw)

is_t = b1.apply(is_thrush)

# 3회 모두 1순위 일치 여부
all_same = (b1 == b2) & (b2 == b3)
# 적어도 2회 일치
two_same = ((b1 == b2) | (b2 == b3) | (b1 == b3))

thrush_total  = is_t.sum()
other_total   = (~is_t).sum()
thrush_all3   = (is_t & all_same).sum()
other_all3    = (~is_t & all_same).sum()
thrush_two    = (is_t & two_same).sum()
other_two     = (~is_t & two_same).sum()

print(f"지빠귀류 샘플: {thrush_total}개")
print(f"  3회 모두 일치: {thrush_all3}개 ({thrush_all3/thrush_total*100:.1f}%)")
print(f"  2회 이상 일치: {thrush_two}개 ({thrush_two/thrush_total*100:.1f}%)")
print(f"다른 새 샘플: {other_total}개")
print(f"  3회 모두 일치: {other_all3}개 ({other_all3/other_total*100:.1f}%)")
print(f"  2회 이상 일치: {other_two}개 ({other_two/other_total*100:.1f}%)")

# ── 그림 8: BirdNET 재현성 차트 생성 ─────────────────────────────
p_t3 = round(thrush_all3/thrush_total*100, 1) if thrush_total > 0 else 0
p_o3 = round(other_all3/other_total*100, 1) if other_total > 0 else 0

fig, ax = plt.subplots(figsize=(6.5, 4.5), facecolor='white')
cats = [f'지빠귀류\n(n={thrush_total})', f'다른 새\n(n={other_total})']
vals = [p_t3, p_o3]
colors = [RED, GREEN_DARK]
bars = ax.bar(cats, vals, color=colors, width=0.45, edgecolor='white', linewidth=1.5)
for bar, val in zip(bars, vals):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1.5,
            f'{val}%', ha='center', va='bottom', fontsize=16, fontweight='bold',
            color=bar.get_facecolor())
ax.set_ylabel('3회 반복 시 1순위 일치율 (%)', fontsize=11)
ax.set_title('BirdNET 같은 소리 3회 반복 테스트 결과\n(같은 결과가 나오는 비율)', fontsize=12, fontweight='bold', pad=12)
ax.set_ylim(0, 110)
ax.set_facecolor(BG)
ax.grid(axis='y', alpha=0.3, linestyle='--')
ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
ax.tick_params(labelsize=12)

# 해석 텍스트 박스
desc = f'지빠귀류는 3번 반복해도\n{p_t3}%만 같은 결과 → 일관성 낮음'
ax.text(0, p_t3 / 2, desc, ha='center', va='center', fontsize=9,
        color='white', fontweight='bold')

chart8_path = os.path.join(CHART_DIR, 'chart8_재현성.png')
fig.savefig(chart8_path, dpi=180, bbox_inches='tight', facecolor='white')
plt.close(fig)
print(f"chart8 저장: {chart8_path}")

# ── 차트 삽입 ──────────────────────────────────────────────────────
CHART_MAP = {
    '그림 1': ('chart1_일치율.png', 5.0),
    '그림 2': ('chart2_후보수.png', 4.5),
    '그림 3': ('chart3_지빠귀구별.png', 4.5),
    '그림 4': ('chart4_직박구리구별.png', 4.5),
    '그림 5': ('chart5_확신도.png', 5.0),
    '그림 6': ('chart6_장소별.png', 5.5),
    '그림 7': ('chart7_날씨별.png', 4.0),
    '그림 8': ('chart8_재현성.png', 5.5),
}

doc = Document(SRC_DOCX)
inserted = 0

for para in doc.paragraphs:
    text = para.text.strip()
    if '(그래프 위치)' not in text:
        continue
    for key, (fname, width) in CHART_MAP.items():
        if key in text:
            chart_path = os.path.join(CHART_DIR, fname)
            if os.path.exists(chart_path):
                para.clear()
                run = para.add_run()
                run.add_picture(chart_path, width=Inches(width))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted += 1
                print(f"삽입: {key} → {fname}")
            else:
                print(f"파일 없음: {chart_path}")
            break

doc.save(OUT_DOCX)
print(f"\n완료! 차트 {inserted}개 삽입")
print(f"저장: {OUT_DOCX}")
