#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아"
CHARTS = os.path.join(BASE, "charts")
OUT = os.path.join(BASE, "배이아_최종탐구보고서_수정제출본.docx")

doc = Document()

def shd(cell, color="D5E8F0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), color)
    s.set(qn("w:val"), "clear")
    tcPr.append(s)

def add_table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, ct in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = str(ct)
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.name = "맑은 고딕"
                if i == 0:
                    p.runs[0].bold = True
                    shd(cell)
    return t

def add_img(fname, w=14):
    p = os.path.join(CHARTS, fname)
    if os.path.exists(p):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(p, width=Cm(w))
    else:
        doc.add_paragraph(f"[차트 없음: {fname}]")

def h1(t):
    p = doc.add_heading(t, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
        run.font.name = "맑은 고딕"
    p.paragraph_format.space_before = Pt(18)

def h2(t):
    p = doc.add_heading(t, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x24, 0x71, 0xa3)
        run.font.name = "맑은 고딕"
    p.paragraph_format.space_before = Pt(12)

def body(t):
    p = doc.add_paragraph(t)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def bul(t):
    p = doc.add_paragraph(t, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "맑은 고딕"

def cap(t):
    p = doc.add_paragraph(t)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.italic = True

# ── 제목
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI 새소리 인식 앱은 왜 지빠귀를 헷갈릴까?")
run.font.size = Pt(22); run.font.bold = True; run.font.name = "맑은 고딕"
run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("— BirdNET과 Bird Identifier 앱 비교 분석 —")
r2.font.size = Pt(14); r2.font.name = "맑은 고딕"
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

h1("1. 탐구 동기")
body("우장산 산책 중 스마트폰 새소리 인식 앱 BirdNET으로 새소리를 녹음해보니, 같은 지빠귀 소리를 넣어도 매번 다른 새로 인식하는 것을 발견했다. 또 다른 앱 Bird Identifier와 결과가 다를 때도 많았다. 'AI 앱은 왜 지빠귀를 자꾸 헷갈릴까?'라는 의문에서 이 탐구를 시작했다.")
doc.add_paragraph()

h1("2. 핵심 관점: 정확도보다 일관성")
body("새소리 인식의 정확성을 판단하려면 전문 조류학자의 검증이 필요하다. 그러나 이 탐구에서는 두 앱이 같은 소리에 같은 결과를 내는지(일치율)와, 같은 소리를 반복 테스트했을 때 결과가 안정적인지(재현성)를 중심으로 분석했다.")
doc.add_paragraph()

h1("3. 선행자료 및 배경지식")
h2("BirdNET (Cornell Lab of Ornithology)")
add_table([["항목","BirdNET의 특징"],["개발","미국 코넬 대학교 조류연구소"],["인식 방식","소리를 멜 스펙트로그램으로 변환 후 CNN 딥러닝 분석"],["인식 종수","전 세계 6,000종 이상"],["위치·날짜 활용","녹음 위치와 날짜를 함께 입력하면 그 지역·계절 출현 확률 반영"],["후보 제시","1순위 외 추가 후보 새를 최대 3개까지 제시"]])
doc.add_paragraph()
h2("Bird Identifier by Sound AI ID")
add_table([["항목","Bird Identifier의 특징"],["인식 방식","소리 + 사진 동시 인식 가능"],["특징","AI 기반 실시간 새 종류 판별"]])
doc.add_paragraph()
h2("지빠귀(Thrush)란?")
body("지빠귀는 소형~중형 산새로, 국내에는 개똥지빠귀·흰배지빠귀·되지빠귀 등 여러 종이 있다. 종마다 울음소리가 복잡하고 비슷하며, 계절·지역에 따라 같은 종도 소리 변이가 크다. 이 탐구에서 A·B 표준 음원은 모두 지빠귀 소리(같은 새 종류)이다.")
doc.add_paragraph()

h1("4. 연구 질문과 가설")
add_table([["번호","연구 질문","가설"],["Q1/H1","AI 앱은 지빠귀류를 다른 새보다 더 자주 헷갈릴까?","지빠귀류는 소리가 복잡해 앱 일치율이 낮을 것이다"],["Q2/H2","장소·날씨 환경이 AI 결과에 영향을 줄까?","소음·바람 등 환경이 나쁠수록 후보 새 수가 많을 것이다"],["Q3/H3","사람도 지빠귀 소리 두 개가 같은 새임을 알아차릴까?","지빠귀 소리가 복잡해 정답률이 낮을 것이다"],["Q4/H4","새에 대한 배경지식이 많을수록 정답률이 높을까?","새를 잘 아는 사람이 정답을 더 잘 맞힐 것이다"]])
doc.add_paragraph()

h1("5. 연구 방법")
h2("앱 비교 실험")
add_table([["항목","내용"],["수집 기간","2026년 5월 ~ 6월"],["수집 장소","우장산, 집 근처 공원, 학교 주변, 아파트단지"],["총 샘플 수","47개 (지빠귀류 16개, 다른 새 31개)"],["테스트 방법","동일 녹음을 BirdNET에 3회, Bird Identifier에 1회 입력"],["기록 항목","1·2·3순위 결과, 장소, 날씨, 시간대"]])
doc.add_paragraph()
h2("사람 설문 조사")
add_table([["항목","내용"],["대상","129명 (초등학생 포함 전 연령)"],["방법","구글 폼 온라인 설문"],["1부 음원","A=지빠귀, B=지빠귀 (같은 종 → 정답: 같은 새)"],["2부 음원","C=직박구리, D=꾀꼬리 (다른 종 → 정답: 다른 새)"],["음원 출처","국립생물자원관·Xeno-canto 표준 음원 (현장 녹음 아님)"],["3부 설문","AI 앱 인식도, AI 결과 신뢰도"]])
doc.add_paragraph()

h1("6. 데이터 분석 결과")
h2("6-1. 두 앱의 1순위 일치율 비교")
body("같은 새소리를 BirdNET과 Bird Identifier에 동시에 입력했을 때, 두 앱의 1순위 결과가 일치하는 비율을 비교했다.")
add_table([["분류","샘플 수","일치 수","일치율"],["지빠귀류 식별","16개","1개","6.2%"],["다른 새 식별","31개","16개","51.6%"],["전체","47개","17개","36.2%"]])
doc.add_paragraph()
add_img("chart1_일치율.png", 14)
cap("그림 1. 두 AI 앱의 1순위 결과 일치율 비교")
body("지빠귀류 일치율(6.2%)이 다른 새(51.6%)보다 크게 낮았다. H1 지지.")
doc.add_paragraph()

h2("6-2. BirdNET 평균 후보 새 수 (AI 혼란 지표)")
body("BirdNET이 제시한 후보 새 수를 AI가 얼마나 헷갈리는지의 지표로 분석했다. 후보가 많을수록 AI가 더 불확실하다는 뜻이다.")
add_table([["분류","평균 후보 수"],["지빠귀류 (n=16)","2.69마리"],["다른 새 (n=31)","1.90마리"]])
doc.add_paragraph()
add_img("chart2_후보수.png", 12)
cap("그림 2. BirdNET 평균 후보 새 수 비교 (높을수록 AI가 더 헷갈림)")
doc.add_paragraph()

h2("6-3. 사람 설문 결과")
h2("응답자 구성")
add_table([["연령대","응답자 수","비율"],["10대","103명","79.8%"],["40~50대","20명","15.5%"],["60대 이상","4명","3.1%"],["20~30대","2명","1.6%"]])
doc.add_paragraph()
add_table([["새 지식 수준","응답자 수","비율"],["잘 모른다","71명","55.0%"],["조금 안다","56명","43.4%"],["잘 안다","2명","1.6%"]])
doc.add_paragraph()

h2("A·B 구별 결과 (지빠귀 A vs 지빠귀 B → 정답: 같은 새)")
body("A와 B는 모두 지빠귀 소리다. 같은 새 소리라고 맞힌 비율은 30.2%에 불과했다. 62.0%의 응답자는 다른 새 소리라고 잘못 답했다.")
add_table([["응답","응답자 수","비율","정오"],["같은 새 소리다 (정답)","39명","30.2%","정답"],["다른 새 소리다 (오답)","80명","62.0%","오답"],["모르겠다","10명","7.8%","-"]])
doc.add_paragraph()
add_img("chart3_지빠귀구별.png", 14)
cap("그림 3. 지빠귀 A·B 표준 음원 구별 결과 (정답: 같은 새)")
doc.add_paragraph()

h2("C·D 구별 결과 (직박구리 C vs 꾀꼬리 D → 정답: 다른 새)")
body("C는 직박구리, D는 꾀꼬리 소리다. 다른 새 소리라고 맞힌 비율은 57.4%였다.")
add_table([["응답","응답자 수","비율","정오"],["다른 새 소리다 (정답)","74명","57.4%","정답"],["같은 새 소리다 (오답)","31명","24.0%","오답"],["모르겠다","24명","18.6%","-"]])
doc.add_paragraph()
add_img("chart4_꾀꼬리구별.png", 14)
cap("그림 4. 직박구리·꾀꼬리 C·D 표준 음원 구별 결과 (정답: 다른 새)")
doc.add_paragraph()

h2("확신도 분포 비교")
body("지빠귀 A·B 평균 확신도(3.34점)가 직박구리·꾀꼬리 C·D(3.05점)보다 높았다. 그러나 A·B 정답률(30.2%)이 훨씬 낮아, 사람들이 자신 있게 틀린 경우가 많았다.")
add_table([["점수","A·B 응답자","A·B 비율","C·D 응답자","C·D 비율"],["1점 (전혀 확신 없음)","7명","5.4%","17명","13.2%"],["2점","18명","14.0%","25명","19.4%"],["3점 (보통)","54명","41.9%","47명","36.4%"],["4점","24명","18.6%","14명","10.9%"],["5점 (매우 확신)","26명","20.2%","26명","20.2%"],["평균","3.34점","","3.05점",""]])
doc.add_paragraph()
add_img("chart5_확신도.png", 14)
cap("그림 5. 새소리 구별 평균 확신도 비교")
doc.add_paragraph()

h2("★ 새 지식 수준별 A·B 정답률 교차분석")
body("새에 대한 지식이 많을수록 지빠귀 A·B가 같은 새임을 더 잘 알아차렸다. 잘 안다고 한 응답자(2명)는 100% 정답이었지만, 잘 모른다고 한 응답자(71명)의 정답률은 22.5%에 불과했다.")
add_table([["새 지식 수준","응답자 수","정답(같은 새)","오답+모름","정답률"],["잘 안다","2명","2명","0명","100.0% (샘플 작아 참고)"],["조금 안다","56명","21명","35명","37.5%"],["잘 모른다","71명","16명","55명","22.5%"],["전체","129명","39명","90명","30.2%"]])
doc.add_paragraph()
add_img("chart_교차분석.png", 15)
cap("그림 6. 새 지식 수준별 A·B 정답률 (H4 검증)")
body("새를 잘 알수록 정답률이 높았다(H4 지지). 그러나 전체 정답률 30.2%로, 배경지식이 있어도 지빠귀 소리 구별은 어렵다.")
doc.add_paragraph()

h2("6-4. 환경 변수 분석 (장소·날씨)")
add_table([["장소","샘플 수","BirdNET 평균 후보 수"],["아파트단지","5개","1.40마리"],["학교","4개","1.67마리"],["집 앞 공원","18개","2.17마리"],["우장산 (산)","20개","2.50마리"]])
doc.add_paragraph()
add_table([["날씨","평균 후보 수"],["맑음","1.73마리"],["구름·바람","2.43마리"]])
doc.add_paragraph()
add_img("chart6_장소별.png", 15)
cap("그림 7. 장소별 BirdNET 평균 후보 새 수")
add_img("chart7_날씨별.png", 11)
cap("그림 8. 날씨별 BirdNET 평균 후보 새 수")
body("자연 환경(우장산)과 바람이 많은 날에 후보 새가 더 많이 제시됐다. H2 지지.")
doc.add_paragraph()

h2("6-5. BirdNET 재현성 분석 (같은 소리 3회 반복)")
body("같은 새소리를 BirdNET에 3번 입력했을 때 1순위 결과가 모두 같은 비율(재현성)을 분석했다.")
add_table([["분류","샘플 수","3회 모두 동일","재현율"],["지빠귀류","9개","0개","0.0%"],["다른 새","39개","3개","7.7%"]])
body("지빠귀류는 3번 중 1번도 같은 결과가 나오지 않았다(0.0%). AI에게도 지빠귀가 얼마나 어려운 새인지를 잘 보여준다.")
add_img("chart8_재현성.png", 15)
cap("그림 9. BirdNET 재현성 — 같은 소리 3회 반복 시 1순위 일치율")
doc.add_paragraph()

h2("6-6. AI 앱 인식도 및 신뢰도")
add_table([["AI 앱 인식도","응답자 수","비율"],["처음 알았다","102명","79.1%"],["이미 알고 있었다","23명","17.8%"],["사용 경험 있다","4명","3.1%"]])
doc.add_paragraph()
add_table([["AI 결과 신뢰도","응답자 수","비율"],["매우 신뢰한다","7명","5.4%"],["어느 정도 신뢰한다","70명","54.3%"],["잘 모르겠다","33명","25.6%"],["별로 믿지 않는다","19명","14.7%"],["신뢰 합계","77명","59.7%"]])
doc.add_paragraph()
add_img("chart8_인식도.png", 11)
cap("그림 10. AI 새소리 앱 기능 인식도")
add_img("chart9_신뢰도.png", 15)
cap("그림 11. AI 결과 신뢰도")
doc.add_paragraph()

h1("7. 종합 분석")
add_table([["가설","결과"],["H1: 지빠귀류는 AI 앱 간 일치율이 낮을 것이다","지지 — 지빠귀 6.2% vs 다른 새 51.6%"],["H2: 환경이 나쁠수록 후보 새 수가 많을 것이다","지지 — 우장산 2.50 > 아파트 1.40 / 바람 2.43 > 맑음 1.73"],["H3: 사람도 지빠귀 소리 구별이 어려울 것이다","강하게 지지 — A·B 정답률 30.2% (10명 중 3명만 맞힘)"],["H4: 새 지식이 많을수록 정답률이 높을 것이다","지지 — 잘 안다 100% > 조금 안다 37.5% > 잘 모른다 22.5%"]])
doc.add_paragraph()
body("지빠귀 소리는 AI와 사람 모두에게 어렵다. 사람의 A·B 정답률(30.2%)이 매우 낮아, AI가 헷갈리는 것이 당연하다는 것을 알 수 있었다. 지빠귀류는 BirdNET 3회 반복 테스트에서 0%의 재현성을 보여, AI의 일관성도 매우 낮았다.")
doc.add_paragraph()

h1("8. 결론")
body("이 탐구를 통해 다음 사실을 알게 됐다:")
bul("지빠귀 소리는 AI 앱 간 일치율이 6.2%로, 다른 새(51.6%)보다 훨씬 낮다.")
bul("BirdNET은 지빠귀 후보 새를 평균 2.69개 제시해, 다른 새(1.90개)보다 더 많이 헷갈린다.")
bul("사람도 지빠귀 A·B가 같은 새임을 맞힌 비율이 30.2%로, 10명 중 3명만 정답이었다.")
bul("새를 잘 아는 사람(100%)이 잘 모르는 사람(22.5%)보다 정답률이 훨씬 높아, 배경지식이 중요하다.")
bul("자연 환경(우장산, 바람 부는 날)에서 AI가 더 많이 혼란스러워했다.")
doc.add_paragraph()
body("결론: AI 새소리 인식 앱이 지빠귀를 헷갈리는 이유는, 지빠귀 소리 자체가 종류도 많고 복잡해서 AI뿐 아니라 사람도 구별하기 어려운 새소리이기 때문이다.")
doc.add_paragraph()

h1("9. 한계점과 후속 탐구")
bul("응답자 중 10대가 79.8%로 대부분이었다. 다양한 연령대 응답을 더 많이 수집하면 연령별 차이를 정확히 분석할 수 있다.")
bul("잘 안다고 응답한 사람이 2명뿐이어서 정확한 비교가 어렵다.")
bul("Bird Identifier는 2·3차 반복 테스트 데이터가 부족해 재현성 분석에 포함하지 못했다.")
bul("전문 조류학자의 검증 없이는 AI 결과의 실제 정확도를 알 수 없다.")
bul("후속 탐구: 지빠귀 종별(개똥지빠귀·흰배지빠귀 등) 일치율 비교, 계절별 분석")
doc.add_paragraph()

h1("10. 배우고 느낀 점")
body("처음에는 AI가 틀린다면 앱이 나쁜 것이라고 생각했다. 탐구를 하면서 지빠귀 소리가 실제로 매우 복잡하고 다양해서 전문가도 구별하기 어렵다는 것을 알게 됐다. 사람의 설문 정답률이 30.2%밖에 되지 않는다는 결과가 가장 놀라웠다. AI를 무조건 믿기보다 어떤 원리로 작동하는지 이해하고 사용하는 것이 중요하다는 것을 배웠다.")
doc.add_paragraph()

h1("11. 참고 자료")
bul("BirdNET 공식 사이트 (Cornell Lab of Ornithology): https://birdnet.cornell.edu")
bul("국립생물자원관 생물다양성정보 (표준 새소리 음원): https://species.nibr.go.kr")
bul("Xeno-canto 새소리 데이터베이스: https://xeno-canto.org")

doc.save(OUT)
print("Done:", OUT)
import os
print("Size:", round(os.path.getsize(OUT)/1024, 1), "KB")
