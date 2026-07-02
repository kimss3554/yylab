#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os, copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC  = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아\배이아_최종탐구보고서 (2).docx"
OUT  = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아\배이아_최종탐구보고서_최종완성본.docx"
CHART_DIR = r"C:\Users\yylab\OneDrive\바탕 화면\wiki\wiki\outputs\주제탐구보고서\통계대회\배이아\charts"

# 플레이스홀더 텍스트 → 차트 파일 매핑
CHART_MAP = {
    "그림 1. 두 AI 앱의 1순위 결과 일치율 비교":      ("chart1_일치율.png",    5.0),
    "그림 2. BirdNET 평균 후보 새 수 비교":           ("chart2_후보수.png",    4.5),
    "그림 3. 지빠귀 A·B 표준 음원 구별 결과":         ("chart3_지빠귀구별.png", 4.5),
    "그림 4. 직박구리·까치 C·D 표준 음원 구별 결과":  ("chart4_직박구리구별.png",4.5),
    "그림 5. 새소리 구별 평균 확신도 비교":            ("chart5_확신도.png",    5.0),
    "그림 6. 장소별 평균 후보 새 수":                 ("chart6_장소별.png",    5.5),
    "그림 7. 날씨별 평균 후보 새 수":                 ("chart7_날씨별.png",    4.0),
    "그림 8. AI 새소리 앱 인식도":                   ("chart8_인식도.png",    4.0),
    "그림 9. AI 결과 신뢰도":                        ("chart9_신뢰도.png",    5.5),
}

# 4장, 5장 채울 내용
SECTION_FILL = {
    "4. 연구 질문과 가설": [
        ("heading2", "연구 질문"),
        ("bullet",   "AI 새소리 인식 앱은 지빠귀류를 다른 새보다 더 자주 헷갈릴까?"),
        ("bullet",   "장소와 환경은 AI 결과에 영향을 줄까?"),
        ("bullet",   "사람도 지빠귀 소리를 구별하기 어려울까?"),
        ("heading2", "가설"),
        ("table",    [
            ["가설", "내용"],
            ["H1",  "지빠귀류는 종류가 많고 소리가 비슷해 AI가 구별하기 어려울 것이다"],
            ["H2",  "소음이 많고 새가 많은 환경에서 AI가 더 많은 후보를 제시할 것이다"],
            ["H3",  "사람도 지빠귀류 소리를 완벽하게 구별하지 못할 것이다"],
        ]),
    ],
    "5. 연구 방법": [
        ("heading2", "표본 수집"),
        ("table",    [
            ["항목", "내용"],
            ["수집 기간",  "2026년 5월 ~ 6월"],
            ["수집 장소",  "우장산, 집 근처, 등굣길"],
            ["도구",      "스마트폰 + BirdNET 앱"],
            ["샘플 수",   "47개 (22종: 지빠귀류 16개, 다른 새 31개)"],
        ]),
        ("heading2", "앱 비교 분석"),
        ("body",     "동일한 녹음 파일을 BirdNET과 Bird Identifier by Sound AI ID에 각각 넣어 1순위 결과가 일치하는지 비교하였다. 또한 BirdNET이 제시한 후보 새 수를 난이도 지표로 활용하였다."),
        ("heading2", "환경 조건 분석"),
        ("table",    [
            ["변수", "기록 방식"],
            ["장소", "우장산(산) / 집 앞(공원) / 학원 가는 길 / 아파트단지"],
            ["날씨", "맑음 / 구름·바람"],
            ["시간대", "오전 / 오후 / 저녁"],
        ]),
        ("heading2", "사람 설문 조사"),
        ("table",    [
            ["항목", "내용"],
            ["대상",    "129명 (주로 성인)"],
            ["방법",    "구글 폼 + 카카오톡 링크 배포"],
            ["음원",    "국립생물자원관·Xeno-canto 표준 음원 사용 (현장 녹음 아님)"],
            ["1부",     "지빠귀 A·B → 같은 새인가? 확신도 (1~5점)"],
            ["2부",     "직박구리·까치 C·D → 같은 새인가? 확신도 (1~5점)"],
            ["3부",     "AI 기능 인식도, AI 결과 신뢰도"],
        ]),
    ],
}

doc = Document(SRC)

def add_caption(para_after_idx, caption_text, doc):
    """차트 아래 캡션 문단 삽입"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.italic = True
    return p

def make_table(doc, rows_data):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # 헤더 행 굵게
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                from docx.oxml import OxmlElement
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'D5E8F0')
                shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'auto')
                shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')
                tcPr.append(shd)
    return table

# 단락 순회하면서 플레이스홀더 교체
paragraphs = doc.paragraphs
i = 0
inserted_charts = []
inserted_sections = []

while i < len(doc.paragraphs):
    para = doc.paragraphs[i]
    text = para.text.strip()

    # ── 차트 삽입 ──────────────────────────────────────────────
    for placeholder, (chart_file, width_in) in CHART_MAP.items():
        if placeholder in text:
            chart_path = os.path.join(CHART_DIR, chart_file)
            if os.path.exists(chart_path):
                # 플레이스홀더 단락을 이미지 단락으로 교체
                para.clear()
                run = para.add_run()
                run.add_picture(chart_path, width=Inches(width_in))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted_charts.append(placeholder)
            break

    # ── 섹션 내용 채우기 ───────────────────────────────────────
    for section_title, items in SECTION_FILL.items():
        if text == section_title:
            # 이 단락 다음에 내용 삽입 (XML 조작)
            from docx.oxml.ns import qn
            ref_para = para._element

            def insert_after(ref, new_elem):
                ref.addnext(new_elem)

            # 역순으로 삽입해야 순서 유지
            for item_type, item_content in reversed(items):
                if item_type == "body":
                    p = doc.add_paragraph(item_content)
                    p.style = doc.styles['Normal']
                    insert_after(ref_para, p._element)
                elif item_type == "bullet":
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item_content)
                    insert_after(ref_para, p._element)
                elif item_type == "heading2":
                    p = doc.add_paragraph(style='Heading 2')
                    p.add_run(item_content)
                    insert_after(ref_para, p._element)
                elif item_type == "table":
                    tbl = make_table(doc, item_content)
                    insert_after(ref_para, tbl._element)
                    spacer = doc.add_paragraph("")
                    insert_after(ref_para, spacer._element)

            inserted_sections.append(section_title)
            break

    i += 1

# 저장
doc.save(OUT)

print("Complete!")
print("Charts inserted: " + str(len(inserted_charts)))
for c in inserted_charts:
    print("  - " + c)
print("Sections filled: " + str(len(inserted_sections)))
for s in inserted_sections:
    print("  - " + s)
print("Output: " + OUT)
